"""DOCX document parser using python-docx."""

import logging
import re
from pathlib import Path
from typing import List, Optional
from docx import Document
from .base import BaseParser, ParseResult, Section
from .docx_table_converter import get_table_paragraph_positions, table_to_markdown

logger = logging.getLogger(__name__)

# Pattern-based heading detection (fallback when no Heading styles exist).
# Ordered by specificity: sub-subsection before subsection before section.
# Require uppercase first word after number to reject list items like
# "1. Balance checks, transaction history, statement requests".
_NUMBERED_HEADING_PATTERNS = [
    # 1.2.3 Sub-subsection Title  (level 3)
    (re.compile(r'^\d+\.\d+\.\d+\s+[A-Z]\w'), 3),
    # 1.2 Subsection Title  (level 2)
    (re.compile(r'^\d+\.\d+\s+[A-Z]\w'), 2),
    # 1. Section Title  (level 1)
    (re.compile(r'^\d+\.\s+[A-Z]\w'), 1),
]

# Maximum length for a line to be considered a heading (avoid body text false positives)
_MAX_HEADING_LENGTH = 120

# Characters that indicate a line is body text / list item, not a heading
_LIST_ITEM_CHARS = re.compile(r'[,;]')

# Words that start sentences, not headings.
# Headings are noun phrases ("Executive Summary", "Data Handling").
# Sentences start with articles, pronouns, prepositions, or determiners.
_SENTENCE_STARTERS = frozenset({
    'The', 'This', 'These', 'Those', 'That',
    'A', 'An', 'Each', 'Every', 'All',
    'It', 'Its', 'Our', 'We', 'They',
    'When', 'If', 'For', 'In', 'On', 'At',
    'There', 'Here', 'Any', 'Some', 'No',
    'Most', 'Many', 'Several',
})

# Extract the text after the numbered prefix ("3.2 " -> rest of string)
_NUMBER_PREFIX_RE = re.compile(r'^\d+(?:\.\d+)*\.?\s+(.*)')

# XML namespace for Word documents
_WORD_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006}'

# Formatting-based heading detection (Tier 1.5) constants
_MAX_FMT_HEADING_LEN = 80       # Tighter than 120 for numbered headings
_MAX_ALLCAPS_LEN = 60           # Match PDF parser heuristic
_MAX_ALLCAPS_WORDS = 8          # Match PDF parser heuristic
_MIN_FONT_DELTA_PT = 2.0        # Minimum pt difference vs body text
_BOLD_MAJORITY_THRESHOLD = 0.5  # Fraction of chars that must be bold

_TOC_ENTRY_RE = re.compile(r'^\s*\d+(?:\.\d+)*\.?\s+(.+?)\s+(\d+)\s*$')


def _compute_body_font_size(doc) -> Optional[float]:
    """Compute the weighted median font size of body text paragraphs.

    Iterates all paragraphs, skipping empty ones and those with Heading styles.
    Collects (font_size_pt, char_count) per run where the font size is explicit.
    Returns the weighted median, which is robust against heading outliers.
    Returns None if no explicit font sizes are found (document uses theme defaults).
    """
    samples: list[tuple[float, int]] = []  # (size_pt, char_count)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Skip headings — we want body text baseline only
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            continue
        for run in para.runs:
            if run.font.size is not None and len(run.text.strip()) > 0:
                size_pt = run.font.size.pt
                char_count = len(run.text)
                samples.append((size_pt, char_count))

    if not samples:
        return None

    # Weighted median: sort by font size, walk until cumulative weight >= half
    total_chars = sum(c for _, c in samples)
    samples.sort(key=lambda s: s[0])
    cumulative = 0
    half = total_chars / 2.0
    for size_pt, char_count in samples:
        cumulative += char_count
        if cumulative >= half:
            return size_pt
    return samples[-1][0]


def _detect_formatting_heading(para, body_font_size_pt: Optional[float]) -> bool:
    """Detect if a paragraph is a heading based on visual formatting heuristics.

    Three OR'd heuristics with shared rejection filters:
    1. Bold majority — >50% of characters are in bold runs
    2. Larger font size — dominant font >= body_font_size + 2pt
    3. ALL CAPS — text is uppercase or runs have all_caps property

    Rejection filters (applied before heuristics):
    - Text too long (>80 chars)
    - Contains comma/semicolon (list item)
    - First word is a sentence starter (body text sentence)
    """
    text = para.text.strip()
    if not text:
        return False

    # --- Shared rejection filters ---
    if len(text) > _MAX_FMT_HEADING_LEN:
        return False
    if _LIST_ITEM_CHARS.search(text):
        return False
    words = text.split()
    if words and words[0] in _SENTENCE_STARTERS:
        return False

    runs = para.runs
    if not runs:
        return False

    # Compute per-run character counts for weighted checks
    total_chars = sum(len(r.text) for r in runs)
    if total_chars == 0:
        return False

    # --- Heuristic 1: Bold majority ---
    bold_chars = sum(len(r.text) for r in runs if r.bold is True)
    bold_fraction = bold_chars / total_chars

    if bold_fraction > _BOLD_MAJORITY_THRESHOLD:
        # Reject "bold prefix" pattern: first run bold but last run NOT bold.
        # Catches "Important Note: details here" where "Important Note: " is bold.
        # A true heading is bold throughout — its last run is also bold.
        if len(runs) > 1 and runs[0].bold is True and runs[-1].bold is not True:
            pass  # Reject bold-prefix, fall through to other heuristics
        else:
            return True

    # --- Heuristic 2: Larger font size ---
    if body_font_size_pt is not None:
        # Find dominant font size in this paragraph
        size_chars: dict[float, int] = {}
        for run in runs:
            if run.font.size is not None and len(run.text.strip()) > 0:
                sz = run.font.size.pt
                size_chars[sz] = size_chars.get(sz, 0) + len(run.text)
        if size_chars:
            dominant_size = max(size_chars, key=size_chars.get)
            if dominant_size >= body_font_size_pt + _MIN_FONT_DELTA_PT:
                return True

    # --- Heuristic 3: ALL CAPS ---
    word_count = len(words)
    # Check actual text content (user typed in caps)
    if text.isupper() and len(text) <= _MAX_ALLCAPS_LEN and word_count <= _MAX_ALLCAPS_WORDS:
        return True
    # Check Word's visual all_caps property (text stored lowercase, displayed uppercase)
    allcaps_chars = sum(len(r.text) for r in runs if r.font.all_caps is True)
    if allcaps_chars / total_chars > _BOLD_MAJORITY_THRESHOLD:
        if len(text) <= _MAX_ALLCAPS_LEN and word_count <= _MAX_ALLCAPS_WORDS:
            return True

    return False


def _detect_numbered_heading(text: str) -> int:
    """Detect if text is a numbered heading pattern.

    Four-layer rejection filter:
    1. Length > 120 chars → body text
    2. Contains comma/semicolon → list item ("1. Balance checks, transaction history")
    3. First word after number is a sentence starter → sentence ("1. The system validates...")
    4. Must match numbered pattern with uppercase start → heading

    Returns:
        Heading level (1-3) if detected, 0 if not a heading.
    """
    if len(text) > _MAX_HEADING_LENGTH:
        return 0
    # Reject list items: "1. Balance checks, transaction history..."
    if _LIST_ITEM_CHARS.search(text):
        return 0
    # Reject sentences: "1. The system validates tokens on every request"
    m = _NUMBER_PREFIX_RE.match(text)
    if m:
        rest = m.group(1)
        words = rest.split()
        if words and words[0] in _SENTENCE_STARTERS:
            return 0
    for pattern, level in _NUMBERED_HEADING_PATTERNS:
        if pattern.match(text):
            return level
    return 0


def _has_page_break(para) -> bool:
    """Check if a paragraph contains or is preceded by a page break.

    Detects <w:br w:type="page"/> in the paragraph's XML runs and
    <w:pageBreakBefore/> in paragraph properties.
    """
    # Check for pageBreakBefore in paragraph properties
    pPr = para._element.find(f'{_WORD_NS}pPr')
    if pPr is not None:
        if pPr.find(f'{_WORD_NS}pageBreakBefore') is not None:
            return True

    # Check for explicit page breaks in runs
    for run in para._element.findall(f'{_WORD_NS}r'):
        for br in run.findall(f'{_WORD_NS}br'):
            if br.get(f'{_WORD_NS}type') == 'page':
                return True
    return False


def _build_page_map(doc) -> List[int]:
    """Build a list of page numbers indexed by non-empty paragraph position.

    Scans all paragraphs for explicit page breaks (<w:br w:type="page"/>
    and <w:pageBreakBefore/>). If page breaks are found, uses them for
    accurate page tracking. Otherwise returns empty list (caller falls back
    to estimation).
    """
    page_numbers = []
    current_page = 1
    found_any_break = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # Still check for page breaks in empty paragraphs
            if _has_page_break(para):
                current_page += 1
                found_any_break = True
            continue

        if _has_page_break(para):
            current_page += 1
            found_any_break = True

        page_numbers.append(current_page)

    if not found_any_break:
        return []  # No page breaks found — caller should estimate
    return page_numbers


def _estimate_page(paragraph_index: int, paragraphs_per_page: int = 40) -> int:
    """Estimate page number from paragraph index.

    DOCX has no native page concept — pages are rendered dynamically by the
    editor. This provides a reasonable estimate when no explicit page breaks
    exist.

    Limitations:
    - Documents with many short paragraphs (bullet lists) compress more
      content per page, so estimates will undershoot (e.g., 108 paragraphs
      estimates ~3 pages, actual may be 4-5).
    - Documents with tables, images, or wide margins will overshoot.
    - For traceability, the section *name* is the primary identifier;
      page number is a secondary reference that helps orient the reader.
    """
    return (paragraph_index // paragraphs_per_page) + 1


def _normalize_heading_text(text: str) -> str:
    """Collapse whitespace so heading comparisons survive DOCX formatting quirks."""
    return re.sub(r'\s+', ' ', text).strip()


def _parse_toc_entry(text: str) -> tuple[str, int] | None:
    """Parse a DOCX-generated table-of-contents line into (title, page)."""
    match = _TOC_ENTRY_RE.match(_normalize_heading_text(text))
    if not match:
        return None
    title = _normalize_heading_text(match.group(1))
    if not title or len(title) > _MAX_HEADING_LENGTH:
        return None
    return title, int(match.group(2))


class DOCXParser(BaseParser):
    """Parser for DOCX documents."""

    def parse(self, file_path: Path) -> ParseResult:
        """Parse DOCX document. Falls back to raw ZIP/XML extraction on failure."""
        try:
            doc = Document(file_path)
        except Exception as e:
            logger.warning(
                "python-docx rejected %s (%s), attempting ZIP/XML fallback",
                file_path.name, e,
            )
            return self._parse_zip_fallback(file_path)
        return self._parse_standard(file_path, doc)

    def _parse_zip_fallback(self, file_path: Path) -> ParseResult:
        """Extract text from a .docx ZIP archive when python-docx rejects it.

        Opens the file as a ZIP, finds word/document.xml (or any XML with text),
        and extracts text content using stdlib xml.etree.
        """
        import zipfile
        import xml.etree.ElementTree as ET

        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        try:
            with zipfile.ZipFile(file_path) as zf:
                # Try the standard location first
                xml_candidates = ["word/document.xml"]
                # Also try any XML file that might contain word content
                xml_candidates.extend(
                    n for n in zf.namelist()
                    if n.endswith(".xml") and n not in xml_candidates
                )

                paragraphs: list[str] = []
                for candidate in xml_candidates:
                    if candidate not in zf.namelist():
                        continue
                    try:
                        tree = ET.parse(zf.open(candidate))
                        root = tree.getroot()
                        # Extract text from <w:t> elements
                        for t_elem in root.iter(f"{{{ns}}}t"):
                            if t_elem.text:
                                paragraphs.append(t_elem.text)
                    except ET.ParseError:
                        continue

                    if paragraphs:
                        break  # Found text, stop searching

                if not paragraphs:
                    # Last resort: extract all text from all XML files
                    for name in zf.namelist():
                        if name.endswith(".xml"):
                            try:
                                tree = ET.parse(zf.open(name))
                                for elem in tree.iter():
                                    if elem.text and elem.text.strip():
                                        paragraphs.append(elem.text.strip())
                            except ET.ParseError:
                                continue

        except zipfile.BadZipFile:
            return ParseResult(
                filename=file_path.name,
                sections=[Section(title="Document", content="Unable to read file: not a valid DOCX/ZIP archive.")],
                raw_text="",
                page_count=1,
                metadata={"parser": "docx-failed", "error": "BadZipFile"},
            )

        raw_text = "\n".join(paragraphs)
        if not raw_text.strip():
            return ParseResult(
                filename=file_path.name,
                sections=[Section(title="Document", content="No readable text found in document.")],
                raw_text="",
                page_count=1,
                metadata={"parser": "docx-zip-fallback", "error": "no_text_extracted"},
            )

        # Use text parser logic to find sections
        from .text_parser import TextParser
        text_parser = TextParser()
        sections = text_parser._detect_sections(raw_text.splitlines())
        if not sections:
            sections = [Section(title="Document", content=raw_text)]

        logger.info(
            "ZIP/XML fallback extracted %d paragraphs, %d sections from %s",
            len(paragraphs), len(sections), file_path.name,
        )

        return ParseResult(
            filename=file_path.name,
            sections=sections,
            raw_text=raw_text,
            page_count=1,
            metadata={"parser": "docx-zip-fallback", "paragraph_count": len(paragraphs)},
        )

    def _parse_standard(self, file_path: Path, doc) -> ParseResult:
        """Standard python-docx parsing path."""

        # Pre-scan for explicit page breaks (accurate when present)
        page_map = _build_page_map(doc)
        has_page_breaks = len(page_map) > 0

        # Extract tables and their positions for insertion into text flow
        table_positions = get_table_paragraph_positions(doc)
        table_markdowns = []
        for table in doc.tables:
            markdown = table_to_markdown(table)
            if markdown:
                table_markdowns.append(markdown)

        # Build mapping of paragraph_index -> table_markdown
        table_map = {}
        for idx, pos in enumerate(table_positions):
            if idx < len(table_markdowns):
                table_map[pos] = table_markdowns[idx]

        if table_markdowns:
            logger.info(
                "Extracted %d tables from document as Markdown",
                len(table_markdowns)
            )

        sections = []
        raw_text = []
        current_section_title = None
        current_section_content = []
        current_section_page = 1
        paragraph_index = 0

        for para in doc.paragraphs:
            # Check if there's a table at current paragraph position
            if paragraph_index in table_map:
                table_md = table_map[paragraph_index]
                raw_text.append(table_md)
                current_section_content.append(table_md)

            text = para.text.strip()
            if not text:
                paragraph_index += 1
                continue

            raw_text.append(text)

            # Resolve page number: prefer explicit breaks, fall back to estimate
            if has_page_breaks and paragraph_index < len(page_map):
                current_page = page_map[paragraph_index]
            else:
                current_page = _estimate_page(paragraph_index)

            # Primary: detect headings by Word style
            is_heading = False
            if para.style and para.style.name and para.style.name.startswith('Heading'):
                is_heading = True

            if is_heading:
                # Save previous section
                if current_section_title:
                    sections.append(Section(
                        title=current_section_title,
                        content='\n'.join(current_section_content),
                        page_start=current_section_page,
                    ))
                current_section_title = text
                current_section_content = []
                current_section_page = current_page
            else:
                current_section_content.append(text)

            paragraph_index += 1

        # Check for any remaining tables after the last paragraph
        if paragraph_index in table_map:
            table_md = table_map[paragraph_index]
            raw_text.append(table_md)
            current_section_content.append(table_md)

        # Save last section
        if current_section_title:
            sections.append(Section(
                title=current_section_title,
                content='\n'.join(current_section_content),
                page_start=current_section_page,
            ))

        # Fallback chain: styles -> formatting -> patterns -> single
        detection_method = "style"

        if not sections:
            sections = self._parse_with_formatting(doc, page_map)
            detection_method = "formatting" if sections else detection_method

        if not sections:
            sections = self._parse_with_patterns(raw_text, page_map)
            detection_method = "pattern" if sections else detection_method

        if not sections:
            sections.append(Section(
                title="Document",
                content='\n'.join(raw_text),
                page_start=1,
            ))
            detection_method = "single"

        sections, refined = self._refine_sections_from_toc(sections)
        if refined:
            detection_method = f"{detection_method}+toc"

        max_page = page_map[-1] if page_map else (
            _estimate_page(paragraph_index - 1) if paragraph_index else 1
        )

        return ParseResult(
            filename=file_path.name,
            sections=sections,
            raw_text='\n'.join(raw_text),
            page_count=max(len(doc.sections), max_page),
            metadata={
                "parser": "python-docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(table_markdowns),
                "section_detection": detection_method,
                "page_source": "explicit_breaks" if has_page_breaks else "estimated",
            }
        )

    def _refine_sections_from_toc(
        self,
        sections: List[Section],
    ) -> tuple[List[Section], bool]:
        """Split oversized DOCX sections using table-of-contents entries.

        Some Word documents use heading styles for front matter only, then place
        the actual body under a single heading such as "INHOUDSOPGAVE". When a
        section starts with a TOC block, use those TOC titles to split the body
        back into canonical sections.
        """
        refined_sections: list[Section] = []
        refined_any = False

        for section in sections:
            split_sections = self._split_section_using_toc(section)
            if len(split_sections) > 1:
                refined_any = True
            refined_sections.extend(split_sections)

        return refined_sections, refined_any

    def _split_section_using_toc(self, section: Section) -> List[Section]:
        """Split a section when it contains an embedded table of contents."""
        lines = [line.strip() for line in section.content.splitlines() if line.strip()]
        if len(lines) < 4:
            return [section]

        toc_entries: list[tuple[str, int]] = []
        body_start = 0
        for idx, line in enumerate(lines):
            entry = _parse_toc_entry(line)
            if entry is None:
                body_start = idx
                break
            toc_entries.append(entry)
        else:
            body_start = len(lines)

        if len(toc_entries) < 3 or body_start >= len(lines):
            return [section]

        title_to_page: dict[str, int] = {}
        for title, page in toc_entries:
            title_to_page.setdefault(title, page)

        toc_content = lines[:body_start]
        body_lines = lines[body_start:]
        split_body: list[Section] = []
        current_title: str | None = None
        current_content: list[str] = []
        current_page = section.page_start
        preface_lines: list[str] = []

        for line in body_lines:
            normalized_line = _normalize_heading_text(line)
            if normalized_line in title_to_page:
                if current_title is not None:
                    split_body.append(Section(
                        title=current_title,
                        content="\n".join(current_content).strip(),
                        page_start=current_page,
                    ))
                current_title = normalized_line
                current_content = []
                current_page = title_to_page.get(normalized_line, section.page_start or 1)
                continue

            if current_title is None:
                preface_lines.append(line)
            else:
                current_content.append(line)

        if current_title is None:
            return [section]

        split_body.append(Section(
            title=current_title,
            content="\n".join(current_content).strip(),
            page_start=current_page,
        ))

        if not any(s.content for s in split_body):
            return [section]

        refined_sections: list[Section] = []
        toc_lines = list(toc_content)
        if preface_lines:
            toc_lines.extend(preface_lines)
        if toc_lines:
            refined_sections.append(Section(
                title=section.title,
                content="\n".join(toc_lines).strip(),
                page_start=section.page_start,
                page_end=section.page_end,
            ))
        refined_sections.extend(split_body)

        logger.info(
            "Refined DOCX section '%s' into %d sections using TOC entries",
            section.title,
            len(refined_sections),
        )
        return refined_sections

    def _parse_with_formatting(
        self, doc, page_map: List[int]
    ) -> List[Section]:
        """Detect sections by visual formatting heuristics (Tier 1.5).

        Operates on paragraph objects (needs run-level access for bold/font/caps).
        Called after Heading-style detection fails but before numbered-pattern
        detection, catching documents that use bold text, larger fonts, or
        ALL CAPS for headings without Word styles.
        """
        body_font_size = _compute_body_font_size(doc)
        has_page_breaks = len(page_map) > 0

        sections = []
        current_section_title = None
        current_section_content = []
        current_section_page = 1
        paragraph_index = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Resolve page number
            if has_page_breaks and paragraph_index < len(page_map):
                page = page_map[paragraph_index]
            else:
                page = _estimate_page(paragraph_index)

            if _detect_formatting_heading(para, body_font_size):
                # Save previous section
                if current_section_title:
                    sections.append(Section(
                        title=current_section_title,
                        content='\n'.join(current_section_content),
                        page_start=current_section_page,
                    ))
                current_section_title = text
                current_section_content = []
                current_section_page = page
            else:
                current_section_content.append(text)

            paragraph_index += 1

        # Save last section
        if current_section_title:
            sections.append(Section(
                title=current_section_title,
                content='\n'.join(current_section_content),
                page_start=current_section_page,
            ))

        if sections:
            logger.info(
                "Formatting-based section detection found %d sections "
                "(body_font_size=%.1fpt)",
                len(sections),
                body_font_size or 0.0,
            )

        return sections

    def _parse_with_patterns(
        self, paragraphs: List[str], page_map: List[int]
    ) -> List[Section]:
        """Detect sections by numbered heading patterns in paragraph text.

        Fallback when the document has no Heading styles (all paragraphs are Normal).
        Matches patterns like "1. Executive Summary", "3.2 Authentication".
        Rejects list items like "1. Balance checks, transaction history...".
        """
        has_page_breaks = len(page_map) > 0
        sections = []
        current_section_title = None
        current_section_content = []
        current_section_page = 1

        for idx, text in enumerate(paragraphs):
            level = _detect_numbered_heading(text)

            # Resolve page for this paragraph
            if has_page_breaks and idx < len(page_map):
                page = page_map[idx]
            else:
                page = _estimate_page(idx)

            if level > 0:
                # Save previous section
                if current_section_title:
                    sections.append(Section(
                        title=current_section_title,
                        content='\n'.join(current_section_content),
                        page_start=current_section_page,
                    ))
                current_section_title = text
                current_section_content = []
                current_section_page = page
            else:
                current_section_content.append(text)

        # Save last section
        if current_section_title:
            sections.append(Section(
                title=current_section_title,
                content='\n'.join(current_section_content),
                page_start=current_section_page,
            ))

        if sections:
            logger.info(
                "Pattern-based section detection found %d sections (no Heading styles)",
                len(sections),
            )

        return sections
