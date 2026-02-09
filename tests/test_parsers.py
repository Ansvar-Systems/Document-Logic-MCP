"""Tests for document parsers."""

import tempfile
import pytest
from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Pt
from document_logic_mcp.parsers.pdf_parser import PDFParser
from document_logic_mcp.parsers.docx_parser import (
    DOCXParser,
    _compute_body_font_size,
    _detect_formatting_heading,
)


# ---------------------------------------------------------------------------
# Helper: save in-memory DOCX to a temp file and parse with DOCXParser
# ---------------------------------------------------------------------------

def _parse_docx(doc: DocxDocument) -> "ParseResult":
    """Save an in-memory python-docx Document to a temp file and parse it."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f)
        tmp_path = Path(f.name)
    parser = DOCXParser()
    return parser.parse(tmp_path)


def _make_body_paragraph(doc, text, font_size_pt=11):
    """Add a Normal paragraph with explicit font size."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size_pt)
    return para


def _make_bold_paragraph(doc, text, font_size_pt=11):
    """Add a fully bold paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size_pt)
    return para


# ---------------------------------------------------------------------------
# Existing tests
# ---------------------------------------------------------------------------

@pytest.mark.skipif(
    not (Path(__file__).parent / "fixtures" / "sample.pdf").exists(),
    reason="Test fixture not available"
)
def test_pdf_parser_extracts_text():
    """Test PDF parser extracts text and metadata."""
    pdf_path = Path(__file__).parent / "fixtures" / "sample.pdf"

    parser = PDFParser()
    result = parser.parse(pdf_path)

    assert result.filename == "sample.pdf"
    assert len(result.sections) > 0
    assert result.page_count > 0
    assert len(result.raw_text) > 0


def test_pdf_parser_instantiates():
    """Test PDF parser can be instantiated."""
    parser = PDFParser()
    assert parser is not None


def test_docx_parser_instantiates():
    """Test DOCX parser can be instantiated."""
    parser = DOCXParser()
    assert parser is not None


# ---------------------------------------------------------------------------
# Tier 1.5 Formatting-based heading detection tests
# ---------------------------------------------------------------------------

class TestFormattingBoldHeadings:
    """Test bold-based heading detection (Heuristic 1)."""

    def test_bold_headings_detected(self):
        """Bold paragraphs shorter than 80 chars are detected as headings."""
        doc = DocxDocument()
        _make_bold_paragraph(doc, "Executive Summary")
        _make_body_paragraph(doc, "This report covers the threat analysis.")
        _make_body_paragraph(doc, "Key findings are presented below.")
        _make_bold_paragraph(doc, "Authentication Architecture")
        _make_body_paragraph(doc, "The system uses OAuth 2.0 for authentication.")

        result = _parse_docx(doc)

        assert len(result.sections) == 2
        assert result.sections[0].title == "Executive Summary"
        assert result.sections[1].title == "Authentication Architecture"

    def test_long_bold_paragraph_rejected(self):
        """Bold paragraphs longer than 80 chars are NOT treated as headings."""
        doc = DocxDocument()
        long_text = "This is a very long bold paragraph that should not be treated as a heading because it exceeds the maximum length"
        assert len(long_text) > 80
        _make_bold_paragraph(doc, long_text)
        _make_body_paragraph(doc, "Body text follows.")

        result = _parse_docx(doc)

        # Should fall through to single "Document" since no valid headings
        assert len(result.sections) == 1
        assert result.sections[0].title == "Document"

    def test_bold_prefix_pattern_rejected(self):
        """'Label: explanation' where bold prefix >50% chars but last run not bold."""
        doc = DocxDocument()
        # Bold prefix is >50% of chars, but last run is NOT bold -> reject
        # "Important Security Note: details" = 32 chars
        #  bold portion: 25 chars (78%), non-bold: 7 chars
        para = doc.add_paragraph()
        bold_run = para.add_run("Important Security Note: ")
        bold_run.bold = True
        bold_run.font.size = Pt(11)
        normal_run = para.add_run("details")
        normal_run.font.size = Pt(11)

        _make_body_paragraph(doc, "Some body text here.")
        _make_body_paragraph(doc, "More body text follows.")

        result = _parse_docx(doc)

        # The bold-prefix paragraph should NOT be a heading
        section_titles = [s.title for s in result.sections]
        assert "Important Security Note: details" not in section_titles

    def test_single_bold_word_accepted(self):
        """A single bold word like 'Conclusion' is accepted as heading."""
        doc = DocxDocument()
        _make_bold_paragraph(doc, "Introduction")
        _make_body_paragraph(doc, "Some introductory text goes here.")
        _make_bold_paragraph(doc, "Conclusion")
        _make_body_paragraph(doc, "Final thoughts and recommendations.")

        result = _parse_docx(doc)

        assert len(result.sections) == 2
        assert result.sections[0].title == "Introduction"
        assert result.sections[1].title == "Conclusion"


class TestFormattingFontSizeHeadings:
    """Test font-size-based heading detection (Heuristic 2)."""

    def test_font_size_headings_detected(self):
        """Paragraphs with font >= body + 2pt are detected as headings."""
        doc = DocxDocument()
        # Heading at 14pt, body at 11pt (delta = 3pt > 2pt threshold)
        para1 = doc.add_paragraph()
        run1 = para1.add_run("System Overview")
        run1.font.size = Pt(14)

        _make_body_paragraph(doc, "The system provides core banking services.")
        _make_body_paragraph(doc, "It handles transactions and user accounts.")

        para2 = doc.add_paragraph()
        run2 = para2.add_run("Security Model")
        run2.font.size = Pt(14)

        _make_body_paragraph(doc, "Role-based access control is enforced.")

        result = _parse_docx(doc)

        assert len(result.sections) == 2
        assert result.sections[0].title == "System Overview"
        assert result.sections[1].title == "Security Model"

    def test_no_explicit_font_sizes_returns_none(self):
        """_compute_body_font_size returns None when no explicit sizes set."""
        doc = DocxDocument()
        # Paragraphs without explicit font sizes (use theme defaults)
        doc.add_paragraph("First paragraph of text.")
        doc.add_paragraph("Second paragraph of text.")

        assert _compute_body_font_size(doc) is None


class TestFormattingAllCapsHeadings:
    """Test ALL CAPS heading detection (Heuristic 3)."""

    def test_allcaps_text_detected(self):
        """ALL CAPS short text is detected as heading."""
        doc = DocxDocument()
        _make_body_paragraph(doc, "EXECUTIVE SUMMARY")
        _make_body_paragraph(doc, "This report covers the full threat analysis.")
        _make_body_paragraph(doc, "Key findings are detailed in subsequent sections.")
        _make_body_paragraph(doc, "THREAT ANALYSIS")
        _make_body_paragraph(doc, "Multiple attack vectors were identified.")

        result = _parse_docx(doc)

        assert len(result.sections) == 2
        assert result.sections[0].title == "EXECUTIVE SUMMARY"
        assert result.sections[1].title == "THREAT ANALYSIS"

    def test_allcaps_font_property_detected(self):
        """Word's all_caps font property is detected as heading."""
        doc = DocxDocument()
        # Paragraph with all_caps property (stored lowercase, displayed uppercase)
        para1 = doc.add_paragraph()
        run1 = para1.add_run("Risk Assessment")
        run1.font.all_caps = True
        run1.font.size = Pt(11)

        _make_body_paragraph(doc, "The following risks were identified during analysis.")
        _make_body_paragraph(doc, "Each risk is rated by likelihood and impact.")

        para2 = doc.add_paragraph()
        run2 = para2.add_run("Mitigation Plan")
        run2.font.all_caps = True
        run2.font.size = Pt(11)

        _make_body_paragraph(doc, "Controls are recommended for each risk.")

        result = _parse_docx(doc)

        assert len(result.sections) == 2
        assert result.sections[0].title == "Risk Assessment"
        assert result.sections[1].title == "Mitigation Plan"


class TestFormattingRejectionFilters:
    """Test shared rejection filters for formatting heuristics."""

    def test_sentence_starter_rejected(self):
        """Bold paragraph starting with sentence starter word is rejected."""
        doc = DocxDocument()
        _make_bold_paragraph(doc, "The system overview describes the architecture")
        _make_body_paragraph(doc, "Body text here.")

        result = _parse_docx(doc)

        section_titles = [s.title for s in result.sections]
        assert "The system overview describes the architecture" not in section_titles

    def test_comma_in_bold_rejected(self):
        """Bold paragraph containing comma is rejected (list item)."""
        doc = DocxDocument()
        _make_bold_paragraph(doc, "Authentication, Authorization")
        _make_body_paragraph(doc, "Body text here.")

        result = _parse_docx(doc)

        section_titles = [s.title for s in result.sections]
        assert "Authentication, Authorization" not in section_titles


class TestFormattingTierOrdering:
    """Test correct ordering of detection tiers."""

    def test_word_heading_styles_win(self):
        """Documents with Word Heading styles use style detection, not formatting."""
        doc = DocxDocument()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("Some intro text.")
        doc.add_heading("Methodology", level=1)
        doc.add_paragraph("How we did it.")

        result = _parse_docx(doc)

        assert result.metadata["section_detection"] == "style"
        assert len(result.sections) == 2

    def test_formatting_before_patterns(self):
        """Formatting detection wins when both formatting and patterns could match."""
        doc = DocxDocument()
        # Bold heading (would match formatting tier)
        _make_bold_paragraph(doc, "Executive Summary")
        _make_body_paragraph(doc, "This report covers the analysis.")
        # Numbered heading (would match pattern tier) - but also bold
        _make_bold_paragraph(doc, "Risk Analysis")
        _make_body_paragraph(doc, "Multiple risks identified.")

        result = _parse_docx(doc)

        # Formatting tier should catch both bold headings first
        assert result.metadata["section_detection"] == "formatting"
        assert len(result.sections) == 2

    def test_detection_method_metadata_formatting(self):
        """Metadata reports 'formatting' when formatting detection used."""
        doc = DocxDocument()
        _make_bold_paragraph(doc, "Overview")
        _make_body_paragraph(doc, "Body text for the overview section.")
        _make_bold_paragraph(doc, "Details")
        _make_body_paragraph(doc, "More detailed information follows.")

        result = _parse_docx(doc)

        assert result.metadata["section_detection"] == "formatting"

    def test_tier3_fallback_preserved(self):
        """Documents with no formatting cues still fall back to single 'Document'."""
        doc = DocxDocument()
        # All plain text, no bold, no size variation, no caps, no heading styles
        doc.add_paragraph("This is a plain paragraph with no formatting.")
        doc.add_paragraph("Another plain paragraph follows here.")
        doc.add_paragraph("And a third plain paragraph as well.")

        result = _parse_docx(doc)

        assert len(result.sections) == 1
        assert result.sections[0].title == "Document"
        assert result.metadata["section_detection"] == "single"
